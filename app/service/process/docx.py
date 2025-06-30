from docx import Document
import pandas as pd
import re
import os
from service.search.extract_service import send_text_chunks_to_gpt

UPLOAD_FOLDER = 'upload'
DOWNLOAD_FOLDER = 'download'


def process_word(filename, number, methods, sections):
    """Создаем рейтинг статей по самым упоминаемым авторам и запрашиваемым методам"""
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    print("start")

    """Получить данные с методами и привести данные к приемлимому виду"""
    methods_check = pd.read_excel(filepath, sheet_name='Methods check')
    methods_check['methods_check'] = methods_check['methods_check'].apply(
        lambda row: re.sub(r'[\[\]\']', '', row)
    )
    print("2")
    """Получить данные с встречаемостью авторов"""
    authors_grade = pd.read_excel(filepath, sheet_name='Authors papers')
    """Получить данные с авторами"""
    authors = pd.read_excel(filepath, sheet_name='Authors')
    '''Получить лист с абстрактами'''
    abstracts = pd.read_excel(filepath, sheet_name='Abstract')
    '''Получить лист с секциями "Материалы и методы"'''
    mat_methods = pd.read_excel(filepath, sheet_name='Methods')
    '''Получить лист с секциями "Результаты"'''
    results = pd.read_excel(filepath, sheet_name='Results')
    '''Получить лист с полными текстами'''
    full = pd.read_excel(filepath, sheet_name='Full')
    print("red_f")
    """
    Сортировка статей по встречаемости их авторов;
    Сортировка статей по пересечению множеств заданных методов и найденных методов в статье
    """
    counts = []
    for row in authors["authors"]:
        row = re.sub(r'[""]', '', row)
        count = 0
        authors_lst = [author.strip() for author in row.split(",")]
        for au in authors_lst:
            num = authors_grade[authors_grade["Автор"] == au.strip()]["Количество упоминаний"].to_list()
            if num:
                count += authors_grade[authors_grade["Автор"] == au.strip()]["Количество упоминаний"].to_list()[0]
        average_count = count / len(authors_lst)
        counts.append(average_count)
    authors["counts"] = counts
    authors = authors.sort_values(by="counts", ascending=False)

    methods_grade = methods_check.copy()
    methods_grade['length_methods'] = methods_grade['methods_check'].apply(
        lambda row: len(set(row.split(", ")).intersection(set(methods)))
    )

    grade_df = methods_grade.merge(authors, on=["title", "doi"])
    titles_remove = []  # Названия статей, у которых не найден ни абстракт, ни полный текст
    for title in grade_df['title']:
        abstract = abstracts[abstracts['title'] == title]['abstract']
        full_text = full[full['title'] == title]['full']
        if abstract.item() == '-' and full_text.item() == '-':
            titles_remove.append(title)

    grade_df = grade_df[~grade_df['title'].isin(titles_remove)]
    sort_df = grade_df.sort_values(
        by=['length_methods', 'counts'], ascending=False)[:number]

    titles_filtered = sort_df["title"].to_list()
    print("sorted_Files")

    """Получаем тексты materials&methods, results для экстракции методик"""
    for i, row in enumerate(mat_methods['methods']):
        response = row + results['results'].iloc[i]
        mat_methods.at[i, 'methodic'] = response
    print("texts")

    """Создаем документ и заполняем данными о найденных статьях"""
    document = Document()
    document.add_heading(
        f'Топ {number} статей с запрашиваемыми методами, авторы которых наиболее упоминаемые по запросу', 1
    )
    for title in titles_filtered:
        document.add_heading(title, level=1)  # title
        doi = authors[authors['title'] == title]['doi'].to_list()
        authors_lst = authors[authors['title'] == title]['authors'].to_list()
        document.add_paragraph(f'DOI: {"".join(doi)}')
        document.add_paragraph(f'Authors: {", ".join(authors_lst)}')
        for section in sections:  # add sections
            document.add_paragraph(section, style='Intense Quote')
            section_df = pd.read_excel(filepath, sheet_name=section)
            section_text = section_df[section_df['title'] == title][section.lower()].to_list()
            document.add_paragraph(section_text)

        document.add_paragraph('Methodic', style='Intense Quote')  # add methodic
        methodic_text = mat_methods[mat_methods['title'] == title]['methodic'].to_list()
        document.add_paragraph(send_text_chunks_to_gpt(methodic_text))

        document.add_page_break()

    download_path = os.path.join(DOWNLOAD_FOLDER, filename)
    document.save(download_path)

    return download_path
